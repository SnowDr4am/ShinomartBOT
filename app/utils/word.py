import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

async def generate_storage_word_document(cell_storage, user_data, worker_data):
    """Генерация аккуратного официального документа (1 лист, RU-месяцы, макет из примера)"""
    try:
        # --- Директория ---
        cell_folder = os.path.join("static", "storage_cells", str(cell_storage.cell_id))
        os.makedirs(cell_folder, exist_ok=True)

        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"storage_{timestamp}.docx"
        full_path = os.path.join(cell_folder, filename)

        # --- Документ и базовые настройки ---
        doc = Document()
        section = doc.sections[0]

        # Поля ~ 1.5 см
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

        # Базовый шрифт
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(11)

        # --- Утилиты ---
        def add_title(text: str):
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = p.add_run(text)
            run.bold = True
            run.font.size = Pt(14)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(8)
            return p

        def add_section_header(text: str):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = True
            r.font.size = Pt(12)
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(4)
            return p

        def add_label_value(label: str, value: str):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            r1 = p.add_run(f"{label}: ")
            r1.bold = True
            p.add_run(value if value else "—")
            return p

        # --- Дата сверху справа ---
        created_dt = cell_storage.created_at if getattr(cell_storage, 'created_at', None) else datetime.now()
        date_str = f"{created_dt.day} {RU_MONTHS_GEN[created_dt.month]} {created_dt.year} г."
        p_date = doc.add_paragraph()
        p_date.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        r_date = p_date.add_run(f"Дата оформления: {date_str}")
        r_date.italic = True
        p_date.paragraph_format.space_after = Pt(4)

        # --- Заголовок (как в примере) ---
        add_title("ДОКУМЕНТ О ПРИЁМЕ ШИН И ДИСКОВ НА ХРАНЕНИЕ")  # в примере именно так :contentReference[oaicite:1]{index=1}

        # --- Шапка организации (без рамок, по центру, две строки) ---
        t_head = doc.add_table(rows=2, cols=1)
        t_head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        _remove_table_borders(t_head)

        org_cell = t_head.cell(0, 0)
        _set_cell_margins(org_cell, top=60, bottom=60)
        p_org = org_cell.paragraphs[0]
        p_org.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r_org = p_org.add_run("ШИНОМАРТ")
        r_org.bold = True
        r_org.font.size = Pt(12)

        addr_cell = t_head.cell(1, 0)
        _set_cell_margins(addr_cell, top=20, bottom=60)
        p_addr = addr_cell.paragraphs[0]
        p_addr.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # адрес как в примере: ул. Правды, 64А
        p_addr.add_run("г. Тюмень, ул. Правды, 64А")  # :contentReference[oaicite:2]{index=2}

        # --- Информация о клиенте ---
        add_section_header("ИНФОРМАЦИЯ О КЛИЕНТЕ")
        if user_data:
            add_label_value("ФИО", f"{getattr(user_data, 'name', '')}".strip())
            add_label_value("Телефон", f"{getattr(user_data, 'mobile_phone', '')}".strip())
        else:
            add_label_value("Телефон", f"{cell_storage.user_id}")

        # --- Информация о сотруднике ---
        add_section_header("ИНФОРМАЦИЯ О СОТРУДНИКЕ")
        worker_name = (getattr(worker_data, 'name', '') or getattr(worker_data, 'fio', '') or "").strip()
        worker_phone = (getattr(worker_data, 'mobile_phone', '') or getattr(worker_data, 'phone', '') or "").strip()
        worker_pos = (getattr(worker_data, 'position', '') or getattr(worker_data, 'role', '') or "").strip()

        add_label_value("ФИО", worker_name)
        if worker_pos:
            add_label_value("Должность", worker_pos)
        add_label_value("Телефон", worker_phone)

        # --- Информация об услуге ---
        add_section_header("ИНФОРМАЦИЯ ОБ УСЛУГЕ")
        st = str(cell_storage.storage_type).lower()
        human_type = "Шины с дисками" if ("rim" in st or "диск" in st) else "Шины"
        add_label_value("Тип хранения", human_type)
        add_label_value("Номер ячейки", f"{cell_storage.cell_id}")
        add_label_value("Описание", f"{cell_storage.description or 'Не указано'}")

        # --- Стоимость (жирно, крупнее) ---
        add_section_header("СТОИМОСТЬ УСЛУГИ")
        p_price = doc.add_paragraph()
        p_price.paragraph_format.space_after = Pt(4)
        r_price = p_price.add_run(f"{cell_storage.price} ₽")
        r_price.bold = True
        r_price.font.size = Pt(14)

        # --- Срок хранения ---
        add_section_header("СРОК ХРАНЕНИЯ")
        smonth = getattr(cell_storage, 'scheduled_month', None)
        if smonth:
            year = smonth.year
            month_name = RU_MONTHS_GEN[smonth.month]
            term_text = f"до {month_name} {year} года"
        else:
            term_text = "не указан"
        add_label_value("Срок", term_text)

        # --- Подписи (без жирных рамок, с ФИО в заголовке) ---
        t_sign = doc.add_table(rows=2, cols=2)
        _remove_table_borders(t_sign)
        # autofit устанавливается автоматически, не нужно явно задавать

        t_sign.cell(0, 0).text = f"Сотрудник{f' ({worker_name})' if worker_name else ''}"
        t_sign.cell(0, 1).text = "Клиент"
        for c in (t_sign.cell(0, 0), t_sign.cell(0, 1)):
            _set_cell_margins(c, top=80, bottom=40)
            p = c.paragraphs[0]
            if p.runs:
                p.runs[0].bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        t_sign.cell(1, 0).text = "_________________________"
        t_sign.cell(1, 1).text = "_________________________"
        for c in (t_sign.cell(1, 0), t_sign.cell(1, 1)):
            _set_cell_margins(c, top=100, bottom=60)
            c.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # --- Блок «ВНИМАНИЮ КЛИЕНТА!» (как в примере) ---
        p_warn_h = doc.add_paragraph()
        p_warn_h.paragraph_format.space_before = Pt(10)
        p_warn_h.paragraph_format.space_after = Pt(4)
        r_warn_h = p_warn_h.add_run("ВНИМАНИЮ КЛИЕНТА!")
        r_warn_h.bold = True
        # Пункты: нумерованный список обычными абзацами
        warn_points = [
            "Настоящий документ подтверждает приём шин и/или дисков на ответственное хранение.",
            "При получении имущества клиент обязан предъявить данный акт.",
            "Если клиент не объявится в течение 2 месяцев после окончания срока хранения, организация имеет полное право утилизировать переданные шины и/или диски без дополнительного уведомления.",
            "Для выдачи имущества клиент обязан предупредить минимум за 2 суток (не считая выходных и праздничных дней).",
        ]  # текст взят из твоего примера :contentReference[oaicite:3]{index=3}
        for i, txt in enumerate(warn_points, start=1):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.add_run(f"{i}. {txt}")

        # --- Единообразные интервалы в документе ---
        for p in doc.paragraphs:
            fmt = p.paragraph_format
            if fmt.space_before is None:
                fmt.space_before = Pt(0)
            if fmt.space_after is None:
                fmt.space_after = Pt(2)
            fmt.line_spacing = 1.0

        # --- Сохранение ---
        doc.save(full_path)
        return full_path

    except ImportError:
        # Фолбэк в .txt (с данными сотрудника)
        cell_folder = os.path.join("static", "storage_cells", str(cell_storage.cell_id))
        os.makedirs(cell_folder, exist_ok=True)
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        filename = f"storage_{timestamp}.txt"
        full_path = os.path.join(cell_folder, filename)

        created_dt = cell_storage.created_at if getattr(cell_storage, 'created_at', None) else datetime.now()
        date_str = f"{created_dt.day} {RU_MONTHS_GEN[created_dt.month]} {created_dt.year} г."
        smonth = getattr(cell_storage, 'scheduled_month', None)
        term_text = f"до {RU_MONTHS_GEN[smonth.month]} {smonth.year} года" if smonth else "не указан"
        st = str(cell_storage.storage_type).lower()
        human_type = "Шины с дисками" if ("rim" in st or "диск" in st) else "Шины"
        worker_name = (getattr(worker_data, 'name', '') or getattr(worker_data, 'fio', '') or "").strip()
        worker_phone = (getattr(worker_data, 'mobile_phone', '') or getattr(worker_data, 'phone', '') or "").strip()
        worker_pos = (getattr(worker_data, 'position', '') or getattr(worker_data, 'role', '') or "").strip()

        with open(full_path, 'w', encoding='utf-8') as f:
            f.write(f"Дата оформления: {date_str}\n\n")
            f.write("ДОКУМЕНТ О ПРИЁМЕ ШИН И ДИСКОВ НА ХРАНЕНИЕ\n")  # :contentReference[oaicite:4]{index=4}
            f.write("ШИНОМАРТ\n")
            f.write("г. Тюмень, ул. Правды, 64А\n")  # :contentReference[oaicite:5]{index=5}
            f.write("\n")
            if user_data:
                f.write(f"ФИО: {getattr(user_data, 'name', '')}\n")
                f.write(f"Телефон: {getattr(user_data, 'mobile_phone', '')}\n")
            else:
                f.write(f"Телефон: {cell_storage.user_id}\n")
            f.write("\nСОТРУДНИК:\n")
            f.write(f"ФИО: {worker_name or '—'}\n")
            if worker_pos:
                f.write(f"Должность: {worker_pos}\n")
            f.write(f"Телефон: {worker_phone or '—'}\n")
            f.write("\n")
            f.write(f"Тип хранения: {human_type}\n")
            f.write(f"Номер ячейки: {cell_storage.cell_id}\n")
            f.write(f"Описание: {cell_storage.description or 'Не указано'}\n")
            f.write(f"Стоимость: {cell_storage.price} ₽\n")
            f.write(f"Срок: {term_text}\n\n")
            f.write("Сотрудник ({}): _________________________\n".format(worker_name if worker_name else ""))
            f.write("Клиент:           _________________________\n\n")
            f.write("ВНИМАНИЮ КЛИЕНТА!\n")
            for i, txt in enumerate(warn_points, start=1):
                f.write(f"{i}. {txt}\n")
        return full_path

    except Exception as e:
        print(f"Ошибка при генерации документа: {e}")
        return None

# ------------------------------
# УТИЛИТЫ
# ------------------------------

RU_MONTHS_GEN = {
    1: "января",  2: "февраля", 3: "марта",
    4: "апреля",  5: "мая",     6: "июня",
    7: "июля",    8: "августа", 9: "сентября",
    10: "октября", 11: "ноября", 12: "декабря",
}

def _set_cell_margins(cell, top=80, start=140, bottom=80, end=140):
    """
    Ставит внутренние отступы в ячейке, предварительно удаляя существующий <w:tcMar>.
    Никакого .xpath — удаляем детей по локальному имени тега.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # удалить существующий <w:tcMar>, если есть
    for child in list(tcPr):
        if child.tag.endswith('tcMar'):
            tcPr.remove(child)

    tcMar = OxmlElement('w:tcMar')
    for tag, val in (('top', top), ('start', start), ('bottom', bottom), ('end', end)):
        node = OxmlElement(f'w:{tag}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)

    tcPr.append(tcMar)


def _remove_table_borders(table):
    """
    Убирает границы у таблицы. Без .xpath — работаем с дочерними элементами напрямую.
    """
    tbl = table._element

    # получить/создать <w:tblPr>
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # удалить существующие <w:tblBorders>, если есть
    for child in list(tblPr):
        if child.tag.endswith('tblBorders'):
            tblPr.remove(child)

    # поставить <w:tblBorders> c val="nil"
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'nil')
        tblBorders.append(el)

    tblPr.append(tblBorders)